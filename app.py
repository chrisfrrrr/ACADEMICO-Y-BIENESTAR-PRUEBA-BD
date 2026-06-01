import io
import re
import zipfile
from pathlib import Path
from datetime import datetime, date
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import plotly.express as px
import requests
import streamlit as st

try:
    from supabase import create_client
except Exception:
    create_client = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
except Exception:
    Document = None

APP_NAME = "Sistema PRO de Seguimiento y Derivación Académica AVE"
RIESGO_ORDEN = {"Bajo": 1, "Moderado": 2, "Alto": 3}
RIESGO_COLOR = {"Bajo": "#0F766E", "Moderado": "#B7791F", "Alto": "#B91C1C"}

SHEETS = {
    "Estudiantes": ["carne", "nombre", "correo", "telefono", "carrera", "trimestre", "curso", "seccion", "canvas_user_id", "asesor_bienestar"],
    "Asesores_Bienestar": ["id_asesor", "nombre", "correo", "telefono", "observaciones"],
    "Historial_Estudiantes": ["fecha", "id_consulta", "carne", "nombre", "correo", "curso", "seccion", "semana", "actividades_pct", "promedio", "entregas_tarde", "semanas_sin_entregas", "ingresos_semana", "dias_inactivo", "horas_respuesta", "riesgo", "riesgo_anterior", "cambio", "motivo_detectado", "asesor_academico"],
    "Derivaciones": ["id_derivacion", "fecha", "carne", "nombre", "correo", "curso", "seccion", "riesgo", "prioridad", "asesor_bienestar", "correo_bienestar", "motivo", "acciones_previas", "observaciones", "estado_derivacion", "asesor_academico"],
    "Mensajes_Enviados": ["fecha", "carne", "nombre", "correo", "curso", "riesgo", "tipo_mensaje", "mensaje_generado", "enviado_canvas", "asesor_academico"],
    "Consultas_Canvas": ["id_consulta", "fecha_consulta", "asesor_academico", "curso", "semana", "total_estudiantes", "bajo", "moderado", "alto", "fuente_datos"],
    "Configuracion": ["parametro", "valor"],
}

SUPABASE_TABLE = "ave_base_datos"


# -----------------------------------------------------------------------------
# UI
# -----------------------------------------------------------------------------
st.set_page_config(page_title="Seguimiento AVE PRO", page_icon="🎓", layout="wide")
st.markdown("""
<style>
.main .block-container{padding-top:1.1rem; padding-bottom:2rem;}
.metric-card{border-radius:18px;padding:18px;background:#ffffff;border:1px solid #e5e7eb;box-shadow:0 8px 22px rgba(15,23,42,.06)}
.small-muted{font-size:0.87rem;color:#64748b}.risk-low{color:#0F766E;font-weight:700}.risk-mid{color:#B7791F;font-weight:700}.risk-high{color:#B91C1C;font-weight:700}
.section-title{font-size:1.2rem;font-weight:800;margin-top:1rem;color:#0F172A}.pill{border-radius:999px;padding:4px 10px;background:#f1f5f9;color:#334155;font-size:.85rem}
</style>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# Helpers DB
# -----------------------------------------------------------------------------
def empty_db() -> Dict[str, pd.DataFrame]:
    return {name: pd.DataFrame(columns=cols) for name, cols in SHEETS.items()}


def normalize_cols(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip().lower().replace(" ", "_").replace("á", "a").replace("é", "e").replace("í", "i").replace("ó", "o").replace("ú", "u") for c in df.columns]
    return df




def clean_key_value(value):
    """Normaliza identificadores para evitar errores de merge por tipos mixtos.

    Excel puede leer el carné como número entero, mientras Canvas lo entrega como texto.
    Esta función convierte ambos casos a texto comparable, preservando vacíos como NaN.
    """
    if pd.isna(value):
        return np.nan
    txt = str(value).strip()
    if txt.lower() in ["", "nan", "none", "nat"]:
        return np.nan
    # Cuando Excel lee 20261234 como 20261234.0, se limpia el decimal artificial.
    if re.fullmatch(r"\d+\.0", txt):
        txt = txt[:-2]
    return txt


def normalize_key_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Convierte claves institucionales y Canvas a texto antes de comparar o fusionar."""
    df = df.copy()
    for col in ["carne", "canvas_user_id", "correo", "login_id", "id_asesor"]:
        if col in df.columns:
            df[col] = df[col].map(clean_key_value)
    return df


def normalize_db_keys(db: Dict[str, pd.DataFrame]) -> Dict[str, pd.DataFrame]:
    """Aplica normalización de claves a todas las hojas de la base."""
    out = {}
    for name, df in db.items():
        out[name] = normalize_key_columns(df) if isinstance(df, pd.DataFrame) else df
    return out

def load_excel_db(uploaded_file) -> Dict[str, pd.DataFrame]:
    db = empty_db()
    if uploaded_file is None:
        return db
    xls = pd.ExcelFile(uploaded_file)
    for sheet in xls.sheet_names:
        key = next((s for s in SHEETS if s.lower() == sheet.lower()), sheet)
        if key in SHEETS:
            df = pd.read_excel(xls, sheet_name=sheet)
            df = normalize_cols(df)
            for col in SHEETS[key]:
                if col not in df.columns:
                    df[col] = np.nan
            db[key] = normalize_key_columns(df[SHEETS[key]])
    return normalize_db_keys(db)


def export_db_excel(db: Dict[str, pd.DataFrame]) -> bytes:
    """Exporta la base a Excel evitando fallos por columnas vacías, None, NaN o tipos mixtos."""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        for name, cols in SHEETS.items():
            df = db.get(name, pd.DataFrame(columns=cols)).copy()
            if name == "Estudiantes" and (df.empty or len(df.dropna(how="all")) == 0):
                df = latest_students_from_history(db)

            # Asegurar estructura mínima de cada hoja
            for c in cols:
                if c not in df.columns:
                    df[c] = ""
            df = df[cols].copy()

            # Reemplazar valores problemáticos antes de exportar
            df = df.replace({np.nan: "", None: "", pd.NaT: ""})
            df.to_excel(writer, index=False, sheet_name=name)

            workbook = writer.book
            ws = writer.sheets[name]
            header_fmt = workbook.add_format({"bold": True, "bg_color": "#0B1F3A", "font_color": "#FFFFFF", "border": 1})
            body_fmt = workbook.add_format({"text_wrap": True, "valign": "top"})

            for i, col in enumerate(df.columns):
                ws.write(0, i, col, header_fmt)

                # Cálculo seguro del ancho de columna. Evita ValueError cuando quantile devuelve NaN.
                try:
                    if len(df) > 0:
                        lengths = df[col].fillna("").astype(str).str.len()
                        q75 = lengths.quantile(0.75)
                        if pd.isna(q75):
                            q75 = len(str(col))
                        base_width = int(max(q75, len(str(col)))) + 3
                    else:
                        base_width = len(str(col)) + 3
                    width = min(max(12, base_width), 36)
                except Exception:
                    width = min(max(12, len(str(col)) + 3), 36)

                ws.set_column(i, i, width, body_fmt)
            ws.freeze_panes(1, 0)
    return output.getvalue()


def append_rows(db: Dict[str, pd.DataFrame], sheet: str, rows: List[Dict]) -> Dict[str, pd.DataFrame]:
    if not rows:
        return db
    new_df = pd.DataFrame(rows)
    for col in SHEETS[sheet]:
        if col not in new_df.columns:
            new_df[col] = np.nan
    base = normalize_key_columns(db.get(sheet, pd.DataFrame(columns=SHEETS[sheet])))
    db[sheet] = normalize_key_columns(pd.concat([base, normalize_key_columns(new_df[SHEETS[sheet]])], ignore_index=True))
    return db


def latest_students_from_history(db: Dict[str, pd.DataFrame]) -> pd.DataFrame:
    """Construye/actualiza la hoja Estudiantes con el último registro conocido de cada estudiante.
    Esto evita que el Excel exportado parezca vacío cuando el análisis ya fue registrado
    en Historial_Estudiantes, pero todavía no existía una base maestra cargada.
    """
    hist = normalize_key_columns(db.get("Historial_Estudiantes", pd.DataFrame()).copy())
    if hist.empty:
        return pd.DataFrame(columns=SHEETS["Estudiantes"])
    for col in SHEETS["Historial_Estudiantes"]:
        if col not in hist.columns:
            hist[col] = np.nan
    hist["fecha_sort"] = pd.to_datetime(hist.get("fecha"), errors="coerce")
    keys = []
    if "carne" in hist.columns and hist["carne"].notna().any():
        keys = ["carne"]
    elif "canvas_user_id" in hist.columns and hist["canvas_user_id"].notna().any():
        keys = ["canvas_user_id"]
    else:
        keys = ["correo"]
    latest = hist.sort_values("fecha_sort").drop_duplicates(keys, keep="last")
    out = pd.DataFrame(columns=SHEETS["Estudiantes"])
    for c in SHEETS["Estudiantes"]:
        if c in latest.columns:
            out[c] = latest[c]
        else:
            out[c] = np.nan
    return out[SHEETS["Estudiantes"]].reset_index(drop=True)


def upsert_students_from_analysis(db: Dict[str, pd.DataFrame], analyzed: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    """Actualiza la hoja Estudiantes con los alumnos del análisis activo sin duplicarlos."""
    if analyzed is None or analyzed.empty:
        return db
    current = normalize_key_columns(db.get("Estudiantes", pd.DataFrame(columns=SHEETS["Estudiantes"])).copy())
    for c in SHEETS["Estudiantes"]:
        if c not in current.columns:
            current[c] = np.nan
    incoming = pd.DataFrame(columns=SHEETS["Estudiantes"])
    for c in SHEETS["Estudiantes"]:
        incoming[c] = analyzed[c] if c in analyzed.columns else np.nan
    incoming = normalize_key_columns(incoming[SHEETS["Estudiantes"]].copy())
    # Clave preferida: carne; si no hay carné, usar canvas_user_id; si no, correo.
    combined = pd.concat([current[SHEETS["Estudiantes"]], incoming], ignore_index=True)
    combined = combined.replace({"": np.nan})
    if combined["carne"].notna().any():
        combined["_key"] = combined["carne"].fillna(combined["canvas_user_id"]).fillna(combined["correo"]).astype(str)
    else:
        combined["_key"] = combined["canvas_user_id"].fillna(combined["correo"]).fillna(combined["nombre"]).astype(str)
    combined = combined[combined["_key"].astype(str).str.strip().ne("")]
    combined = combined.drop_duplicates("_key", keep="last").drop(columns=["_key"])
    db["Estudiantes"] = combined[SHEETS["Estudiantes"]].reset_index(drop=True)
    return db



# -----------------------------------------------------------------------------
# Bienestar: carga, normalización y asignación de asesores
# -----------------------------------------------------------------------------
def normalize_person_name(value) -> str:
    """Normaliza nombres para cruces robustos entre Canvas, Excel y bienestar."""
    if pd.isna(value):
        return ""
    txt = str(value).strip().lower()
    replacements = {
        "á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u", "ü": "u", "ñ": "n",
        "Á": "a", "É": "e", "Í": "i", "Ó": "o", "Ú": "u", "Ü": "u", "Ñ": "n",
    }
    for a, b in replacements.items():
        txt = txt.replace(a, b)
    txt = re.sub(r"[^a-z0-9\s]", " ", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt


def find_col(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    """Encuentra una columna aunque venga con acentos, mayúsculas o espacios."""
    normalized = {c: normalize_cols(pd.DataFrame(columns=[c])).columns[0] for c in df.columns}
    candidate_norm = [normalize_cols(pd.DataFrame(columns=[c])).columns[0] for c in candidates]
    for original, norm in normalized.items():
        if norm in candidate_norm:
            return original
    # búsqueda flexible por contenido del nombre
    for original, norm in normalized.items():
        for cand in candidate_norm:
            if cand in norm or norm in cand:
                return original
    return None


def read_bienestar_file(uploaded_file) -> pd.DataFrame:
    """Lee CSV/XLSX de asignación de bienestar y devuelve columnas estándar."""
    if uploaded_file is None:
        raise ValueError("Cargá un archivo CSV o Excel de asignación de bienestar.")
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        # Intentar UTF-8 y fallback común de Excel en español.
        raw_bytes = uploaded_file.getvalue()
        try:
            df = pd.read_csv(io.BytesIO(raw_bytes), encoding="utf-8-sig")
        except Exception:
            df = pd.read_csv(io.BytesIO(raw_bytes), encoding="latin-1")
    else:
        df = pd.read_excel(uploaded_file)
    if df.empty:
        raise ValueError("El archivo de bienestar está vacío.")

    carne_col = find_col(df, ["carne", "carné", "carnet", "carné estudiante", "id estudiante", "codigo", "código"])
    nombre_col = find_col(df, ["nombre completo", "nombre", "estudiante", "nombre del estudiante", "alumno"])
    asesor_col = find_col(df, ["asesor de bienestar", "asesor bienestar", "bienestar", "asesor", "nombre asesor bienestar"])
    correo_col = find_col(df, ["correo asesor", "correo de asesor", "email asesor", "correo bienestar", "mail asesor"])

    if nombre_col is None:
        raise ValueError("No pude identificar la columna del nombre del estudiante en el archivo de bienestar.")
    if asesor_col is None:
        raise ValueError("No pude identificar la columna del asesor de bienestar en el archivo de bienestar.")

    out = pd.DataFrame()
    out["carne"] = df[carne_col] if carne_col else np.nan
    out["nombre"] = df[nombre_col]
    out["asesor_bienestar"] = df[asesor_col]
    out["correo_bienestar"] = df[correo_col] if correo_col else ""
    out = normalize_key_columns(out)
    out["nombre_key"] = out["nombre"].map(normalize_person_name)
    out["asesor_bienestar"] = out["asesor_bienestar"].astype(str).str.strip()
    out = out[out["asesor_bienestar"].notna() & out["asesor_bienestar"].astype(str).str.strip().ne("")]
    out = out[out["asesor_bienestar"].astype(str).str.lower().ne("nan")]
    out = out.drop_duplicates(subset=["carne", "nombre_key"], keep="last")
    return out[["carne", "nombre", "nombre_key", "asesor_bienestar", "correo_bienestar"]]


def apply_bienestar_assignment(db: Dict[str, pd.DataFrame], bienestar_df: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    """Actualiza Estudiantes y Asesores_Bienestar con el archivo cargado."""
    if bienestar_df is None or bienestar_df.empty:
        return db
    db = normalize_db_keys(db)
    est = db.get("Estudiantes", pd.DataFrame(columns=SHEETS["Estudiantes"])).copy()
    for c in SHEETS["Estudiantes"]:
        if c not in est.columns:
            est[c] = np.nan
    est = normalize_key_columns(est)
    if "nombre_key" not in est.columns:
        est["nombre_key"] = est["nombre"].map(normalize_person_name)

    b = bienestar_df.copy()
    # Cruce por carné cuando existe
    if "carne" in est.columns and est["carne"].notna().any() and b["carne"].notna().any():
        est = est.merge(b[["carne", "asesor_bienestar", "correo_bienestar"]].dropna(subset=["carne"]).drop_duplicates("carne"), on="carne", how="left", suffixes=("", "_bienestar"))
        est["asesor_bienestar"] = est.get("asesor_bienestar_bienestar").combine_first(est.get("asesor_bienestar"))
        est.drop(columns=[c for c in ["asesor_bienestar_bienestar", "correo_bienestar"] if c in est.columns], inplace=True, errors="ignore")
    # Cruce por nombre normalizado para los que siguen sin asesor
    est = est.merge(b[["nombre_key", "asesor_bienestar", "correo_bienestar"]].drop_duplicates("nombre_key"), on="nombre_key", how="left", suffixes=("", "_nombre"))
    est["asesor_bienestar"] = est.get("asesor_bienestar").combine_first(est.get("asesor_bienestar_nombre"))
    est.drop(columns=[c for c in ["asesor_bienestar_nombre", "correo_bienestar", "nombre_key"] if c in est.columns], inplace=True, errors="ignore")
    db["Estudiantes"] = est[SHEETS["Estudiantes"]].reset_index(drop=True)

    asesores = b[["asesor_bienestar", "correo_bienestar"]].drop_duplicates().copy()
    asesores = asesores.rename(columns={"asesor_bienestar": "nombre", "correo_bienestar": "correo"})
    asesores["id_asesor"] = asesores["nombre"].map(lambda x: re.sub(r"[^A-Za-z0-9]+", "_", str(x)).strip("_").lower())
    asesores["telefono"] = ""
    asesores["observaciones"] = "Importado desde archivo de asignación de bienestar"
    for c in SHEETS["Asesores_Bienestar"]:
        if c not in asesores.columns:
            asesores[c] = ""
    current = db.get("Asesores_Bienestar", pd.DataFrame(columns=SHEETS["Asesores_Bienestar"])).copy()
    combined = pd.concat([current, asesores[SHEETS["Asesores_Bienestar"]]], ignore_index=True)
    combined = normalize_key_columns(combined).drop_duplicates(subset=["nombre"], keep="last")
    db["Asesores_Bienestar"] = combined[SHEETS["Asesores_Bienestar"]].reset_index(drop=True)
    return db


def apply_bienestar_to_analysis(analysis_df: pd.DataFrame, db: Dict[str, pd.DataFrame]) -> pd.DataFrame:
    """Completa asesor_bienestar del análisis activo usando la hoja Estudiantes."""
    if analysis_df is None or analysis_df.empty:
        return analysis_df
    est = normalize_key_columns(db.get("Estudiantes", pd.DataFrame()).copy())
    if est.empty or "asesor_bienestar" not in est.columns:
        return analysis_df
    df = normalize_key_columns(analysis_df.copy())
    if "asesor_bienestar" not in df.columns:
        df["asesor_bienestar"] = np.nan
    # por carné
    if "carne" in df.columns and "carne" in est.columns:
        map_carne = est.dropna(subset=["carne"]).drop_duplicates("carne").set_index("carne")["asesor_bienestar"].to_dict()
        df["asesor_bienestar"] = df["asesor_bienestar"].combine_first(df["carne"].map(map_carne))
    # por nombre
    est["nombre_key"] = est["nombre"].map(normalize_person_name)
    df["nombre_key_tmp"] = df["nombre"].map(normalize_person_name)
    map_nombre = est.dropna(subset=["nombre_key"]).drop_duplicates("nombre_key").set_index("nombre_key")["asesor_bienestar"].to_dict()
    df["asesor_bienestar"] = df["asesor_bienestar"].combine_first(df["nombre_key_tmp"].map(map_nombre))
    df.drop(columns=["nombre_key_tmp"], inplace=True, errors="ignore")
    return df

# -----------------------------------------------------------------------------
# Supabase online database helpers
# -----------------------------------------------------------------------------
def clean_records_for_json(df: pd.DataFrame) -> List[Dict]:
    """Convierte un DataFrame a registros JSON seguros para Supabase."""
    if df is None or df.empty:
        return []
    temp = df.copy()
    for col in temp.columns:
        temp[col] = temp[col].map(lambda x: x.isoformat() if hasattr(x, "isoformat") and not isinstance(x, str) else x)
    temp = temp.astype(object).where(pd.notna(temp), None)
    return temp.to_dict(orient="records")


def get_supabase_config() -> dict:
    """Lee credenciales desde Streamlit Secrets cuando existan."""
    try:
        cfg = dict(st.secrets.get("supabase", {}))
    except Exception:
        cfg = {}
    return cfg


def get_supabase_client(url: str, key: str):
    if create_client is None:
        raise RuntimeError("El paquete supabase no está instalado. Ejecutá: pip install supabase")
    if not url or not key:
        raise ValueError("Falta la URL o la key de Supabase. Configuralas en Streamlit Secrets o ingresalas en la app.")
    return create_client(url.strip(), key.strip())


def read_db_from_supabase(url: str, key: str) -> Dict[str, pd.DataFrame]:
    """Lee la base completa desde una tabla JSONB en Supabase."""
    sb = get_supabase_client(url, key)
    db = empty_db()
    try:
        res = sb.table(SUPABASE_TABLE).select("sheet_name,records,updated_at,updated_by").in_("sheet_name", list(SHEETS.keys())).execute()
    except Exception as e:
        raise RuntimeError(f"No se pudo leer Supabase. Verificá que exista la tabla {SUPABASE_TABLE}. Detalle: {e}")
    rows = res.data or []
    for row in rows:
        sheet = row.get("sheet_name")
        if sheet in SHEETS:
            records = row.get("records") or []
            df = pd.DataFrame(records)
            df = normalize_cols(df) if not df.empty else pd.DataFrame(columns=SHEETS[sheet])
            for col in SHEETS[sheet]:
                if col not in df.columns:
                    df[col] = np.nan
            db[sheet] = normalize_key_columns(df[SHEETS[sheet]])
    return normalize_db_keys(db)


def save_db_to_supabase(db: Dict[str, pd.DataFrame], url: str, key: str, updated_by: str = "") -> None:
    """Guarda cada hoja de la base como un JSON en Supabase."""
    sb = get_supabase_client(url, key)
    payloads = []
    now = datetime.utcnow().isoformat() + "Z"
    for sheet, cols in SHEETS.items():
        df = db.get(sheet, pd.DataFrame(columns=cols)).copy()
        if sheet == "Estudiantes" and (df.empty or len(df.dropna(how="all")) == 0):
            df = latest_students_from_history(db)
        for c in cols:
            if c not in df.columns:
                df[c] = np.nan
        df = normalize_key_columns(df[cols])
        payloads.append({
            "sheet_name": sheet,
            "records": clean_records_for_json(df),
            "updated_at": now,
            "updated_by": updated_by or "Streamlit Cloud",
        })
    try:
        sb.table(SUPABASE_TABLE).upsert(payloads, on_conflict="sheet_name").execute()
    except Exception as e:
        raise RuntimeError(f"No se pudo guardar en Supabase: {e}")


def supabase_status(url: str, key: str) -> pd.DataFrame:
    """Devuelve estado de filas por hoja en Supabase."""
    sb = get_supabase_client(url, key)
    res = sb.table(SUPABASE_TABLE).select("sheet_name,records,updated_at,updated_by").execute()
    rows = []
    for item in (res.data or []):
        records = item.get("records") or []
        rows.append({
            "hoja": item.get("sheet_name"),
            "registros": len(records),
            "actualizado": item.get("updated_at"),
            "actualizado_por": item.get("updated_by"),
        })
    return pd.DataFrame(rows).sort_values("hoja") if rows else pd.DataFrame(columns=["hoja", "registros", "actualizado", "actualizado_por"])


# -----------------------------------------------------------------------------
# Canvas Client
# -----------------------------------------------------------------------------
class CanvasClient:
    def __init__(self, base_url: str, token: str):
        self.base_url = base_url.rstrip("/")
        self.headers = {"Authorization": f"Bearer {token}"}

    def _get_paginated(self, path: str, params: Optional[dict] = None) -> List[dict]:
        url = f"{self.base_url}/api/v1{path}"
        params = params or {}
        params.setdefault("per_page", 100)
        out = []
        while url:
            r = requests.get(url, headers=self.headers, params=params, timeout=35)
            r.raise_for_status()
            data = r.json()
            out.extend(data if isinstance(data, list) else [data])
            url = None
            if "next" in r.links:
                url = r.links["next"]["url"]
                params = None
        return out

    def validate(self) -> Tuple[bool, str]:
        try:
            me = requests.get(f"{self.base_url}/api/v1/users/self", headers=self.headers, timeout=20)
            if me.ok:
                data = me.json()
                return True, data.get("name", "Usuario validado")
            return False, f"Canvas respondió {me.status_code}: {me.text[:180]}"
        except Exception as e:
            return False, str(e)

    def courses(self) -> pd.DataFrame:
        """Devuelve los cursos a los que el token tiene acceso.
        La API de Canvas responde únicamente con los cursos visibles para el usuario del token.
        """
        data = self._get_paginated("/courses", {"enrollment_state": "active", "include[]": ["term", "total_students"]})
        rows = []
        for c in data:
            if not c.get("id") or not c.get("name"):
                continue
            term = c.get("term") or {}
            rows.append({
                "id": c.get("id"),
                "name": c.get("name"),
                "course_code": c.get("course_code"),
                "term": term.get("name"),
                "total_students": c.get("total_students"),
                "workflow_state": c.get("workflow_state"),
            })
        df = pd.DataFrame(rows)
        if not df.empty:
            df["label"] = df.apply(lambda r: f"{r.get('name','')} | ID: {r.get('id','')}" + (f" | {r.get('course_code')}" if pd.notna(r.get('course_code')) and str(r.get('course_code')).strip() else ""), axis=1)
        return df

    def users(self, course_id: str) -> pd.DataFrame:
        data = self._get_paginated(f"/courses/{course_id}/users", {"enrollment_type[]": "student"})
        return normalize_key_columns(pd.DataFrame([{"canvas_user_id": u.get("id"), "nombre": u.get("name"), "correo": u.get("email"), "login_id": u.get("login_id")} for u in data]))

    def enrollments(self, course_id: str) -> pd.DataFrame:
        """Obtiene estudiantes con calificación y actividad cuando Canvas lo permite."""
        data = self._get_paginated(
            f"/courses/{course_id}/enrollments",
            {"type[]": "StudentEnrollment", "state[]": "active", "include[]": ["user", "grades"]}
        )
        rows = []
        for e in data:
            u = e.get("user") or {}
            grades = e.get("grades") or {}
            current_score = grades.get("current_score")
            final_score = grades.get("final_score")
            last_activity_at = e.get("last_activity_at")
            rows.append({
                "canvas_user_id": e.get("user_id") or u.get("id"),
                "carne": u.get("sis_user_id") or u.get("integration_id") or u.get("login_id"),
                "nombre": u.get("name"),
                "correo": u.get("email") or u.get("login_id"),
                "login_id": u.get("login_id"),
                "promedio": current_score if current_score is not None else final_score,
                "last_activity_at": last_activity_at,
                "total_activity_time": e.get("total_activity_time"),
            })
        return normalize_key_columns(pd.DataFrame(rows))

    def assignments(self, course_id: str) -> pd.DataFrame:
        data = self._get_paginated(f"/courses/{course_id}/assignments", {"include[]": ["submission"]})
        rows = []
        for a in data:
            rows.append({"assignment_id": a.get("id"), "nombre_actividad": a.get("name"), "puntos": a.get("points_possible") or 0, "due_at": a.get("due_at"), "published": a.get("published")})
        return pd.DataFrame(rows)

    def submissions(self, course_id: str, assignment_id: str) -> pd.DataFrame:
        data = self._get_paginated(f"/courses/{course_id}/assignments/{assignment_id}/submissions", {"include[]": ["user"]})
        rows = []
        for s in data:
            u = s.get("user") or {}
            rows.append({"assignment_id": assignment_id, "canvas_user_id": s.get("user_id"), "nombre": u.get("name"), "submitted_at": s.get("submitted_at"), "late": s.get("late"), "missing": s.get("missing"), "score": s.get("score"), "workflow_state": s.get("workflow_state")})
        return normalize_key_columns(pd.DataFrame(rows))

    def course_metrics(self, course_id: str, course_name: str = "") -> pd.DataFrame:
        """Construye una tabla de seguimiento desde Canvas.
        Intenta leer matrícula, calificaciones, actividades publicadas y entregas.
        Si Canvas no entrega una métrica, se deja como NaN; la lógica de riesgo la tratará como alerta.
        """
        base = normalize_key_columns(self.enrollments(course_id))
        if base.empty:
            base = normalize_key_columns(self.users(course_id))
        if base.empty:
            return base

        # Actividad reciente: Canvas suele entregar last_activity_at en enrollments.
        if "last_activity_at" in base.columns:
            today = pd.Timestamp.now(tz=None)
            la = pd.to_datetime(base["last_activity_at"], errors="coerce", utc=True).dt.tz_convert(None)
            base["dias_inactivo"] = (today.normalize() - la.dt.normalize()).dt.days
            # Estimación conservadora de ingresos semanales cuando no existe conteo real.
            base["ingresos_semana"] = np.select(
                [base["dias_inactivo"].le(2), base["dias_inactivo"].between(3, 6), base["dias_inactivo"].ge(7)],
                [3, 1, 0],
                default=np.nan
            )
        else:
            base["dias_inactivo"] = np.nan
            base["ingresos_semana"] = np.nan

        try:
            acts = self.assignments(course_id)
            if not acts.empty:
                acts = acts[acts.get("published", True).fillna(True)].copy()
                # Se toman actividades vencidas o sin fecha de cierre visible; evita contar tareas futuras contra el estudiante.
                if "due_at" in acts.columns:
                    due = pd.to_datetime(acts["due_at"], errors="coerce", utc=True)
                    acts = acts[due.isna() | (due <= pd.Timestamp.utcnow())]
                assignment_ids = acts["assignment_id"].dropna().astype(str).tolist()
            else:
                assignment_ids = []

            subs_all = []
            progress = st.progress(0, text="Leyendo actividades y entregas desde Canvas...") if assignment_ids else None
            for idx, aid in enumerate(assignment_ids):
                try:
                    subs_all.append(self.submissions(course_id, aid))
                except Exception:
                    # Si una tarea no se puede leer, se continúa con las demás.
                    pass
                if progress:
                    progress.progress((idx + 1) / max(len(assignment_ids), 1), text=f"Leyendo entregas {idx+1}/{len(assignment_ids)}")
            if progress:
                progress.empty()

            total_actividades = len(assignment_ids)
            if subs_all and total_actividades > 0:
                subs = pd.concat(subs_all, ignore_index=True)
                subs["submitted_flag"] = subs["submitted_at"].notna() | subs["workflow_state"].astype(str).isin(["submitted", "graded"])
                subs = normalize_key_columns(subs)
                base = normalize_key_columns(base)
                resumen = subs.groupby("canvas_user_id", dropna=False).agg(
                    actividades_entregadas=("submitted_flag", "sum"),
                    entregas_tarde=("late", lambda x: int(pd.Series(x).fillna(False).sum())),
                    entregas_faltantes=("missing", lambda x: int(pd.Series(x).fillna(False).sum())),
                ).reset_index()
                resumen["actividades_pct"] = (resumen["actividades_entregadas"] / total_actividades * 100).round(2)
                resumen["entregas_tarde"] = resumen["entregas_tarde"] + resumen["entregas_faltantes"]
                resumen = normalize_key_columns(resumen)
                base = base.merge(resumen[["canvas_user_id", "actividades_pct", "entregas_tarde"]], on="canvas_user_id", how="left")
                base["actividades_pct"] = base["actividades_pct"].fillna(0)
                base["entregas_tarde"] = base["entregas_tarde"].fillna(total_actividades)
            else:
                base["actividades_pct"] = np.nan
                base["entregas_tarde"] = np.nan
        except Exception as e:
            st.warning(f"Canvas no permitió calcular actividades/entregas automáticamente: {e}")
            base["actividades_pct"] = np.nan
            base["entregas_tarde"] = np.nan

        base["semanas_sin_entregas"] = np.nan
        base["horas_respuesta"] = np.nan
        base["curso"] = course_name
        return base

    def send_message(self, recipients: List[str], subject: str, body: str) -> Tuple[bool, str]:
        payload = {"recipients[]": recipients, "subject": subject, "body": body, "force_new": True}
        try:
            r = requests.post(f"{self.base_url}/api/v1/conversations", headers=self.headers, data=payload, timeout=25)
            if r.ok:
                return True, "Mensaje enviado por Canvas."
            return False, f"No se pudo enviar. Canvas respondió {r.status_code}: {r.text[:180]}"
        except Exception as e:
            return False, str(e)

# -----------------------------------------------------------------------------
# Risk logic and messages
# -----------------------------------------------------------------------------
def classify_student(row: pd.Series) -> Tuple[str, str]:
    pct = pd.to_numeric(row.get("actividades_pct"), errors="coerce")
    prom = pd.to_numeric(row.get("promedio"), errors="coerce")
    tarde = pd.to_numeric(row.get("entregas_tarde"), errors="coerce")
    sin_ent = pd.to_numeric(row.get("semanas_sin_entregas"), errors="coerce")
    ingresos = pd.to_numeric(row.get("ingresos_semana"), errors="coerce")
    inactivo = pd.to_numeric(row.get("dias_inactivo"), errors="coerce")
    horas_resp = pd.to_numeric(row.get("horas_respuesta"), errors="coerce")
    motivos = []

    high_flags = 0
    mod_flags = 0
    missing_core = 0

    # Criterio crítico: si no existe evidencia de avance académico en Canvas,
    # NO debe clasificarse como bajo. Se asume alerta alta por ausencia de evidencia.
    core_values = [pct, prom, tarde, ingresos, inactivo]
    if all(pd.isna(v) for v in core_values):
        return "Alto", "sin evidencia disponible de avance, calificación, entregas o actividad en Canvas"

    if pd.isna(pct):
        missing_core += 1
        motivos.append("sin dato de porcentaje de actividades")
    else:
        if pct < 50:
            high_flags += 1; motivos.append("bajo cumplimiento de actividades semanales")
        elif pct < 80:
            mod_flags += 1; motivos.append("cumplimiento parcial de actividades")

    if pd.isna(prom):
        missing_core += 1
        motivos.append("sin dato de promedio en Canvas")
    else:
        if prom < 59:
            high_flags += 1; motivos.append("promedio inferior al mínimo esperado")
        elif prom < 70:
            mod_flags += 1; motivos.append("promedio académico en rango de alerta")

    if pd.isna(tarde):
        missing_core += 1
        motivos.append("sin dato de entregas tardías o pendientes")
    else:
        if tarde >= 4:
            high_flags += 1; motivos.append("entregas pendientes o tardías recurrentes")
        elif tarde >= 2:
            mod_flags += 1; motivos.append("entregas irregulares o tardías")

    if pd.notna(sin_ent) and sin_ent >= 2:
        high_flags += 1; motivos.append("sin entregas por dos o más semanas")

    if pd.isna(ingresos):
        missing_core += 1
        motivos.append("sin dato de ingresos semanales a Canvas")
    else:
        if ingresos == 0:
            high_flags += 1; motivos.append("sin ingresos semanales a Canvas")
        elif ingresos <= 2:
            mod_flags += 1; motivos.append("baja frecuencia de ingreso a Canvas")

    if pd.notna(inactivo):
        if inactivo >= 6:
            high_flags += 1; motivos.append("inactividad total en Canvas por seis o más días")
        elif inactivo >= 3:
            mod_flags += 1; motivos.append("actividad limitada en Canvas")

    if pd.notna(horas_resp):
        if horas_resp >= 120:
            high_flags += 1; motivos.append("no responde comunicaciones por cinco o más días")
        elif horas_resp >= 48:
            mod_flags += 1; motivos.append("responde comunicaciones con retraso")

    # Reglas de decisión más estrictas:
    # - Un indicador crítico alto + otra alerta/moderada => Alto.
    # - Dos indicadores críticos altos => Alto.
    # - Varios datos ausentes también suben la alerta, porque no existe evidencia suficiente para bajo.
    if high_flags >= 2 or (high_flags >= 1 and (mod_flags >= 1 or missing_core >= 2)):
        return "Alto", ", ".join(dict.fromkeys(motivos)) or "riesgo académico alto"
    if high_flags == 1:
        return "Moderado", ", ".join(dict.fromkeys(motivos)) or "señales académicas de alerta"
    if mod_flags >= 2 or (mod_flags >= 1 and missing_core >= 1):
        return "Moderado", ", ".join(dict.fromkeys(motivos)) or "señales académicas de alerta"
    if missing_core >= 3:
        return "Alto", ", ".join(dict.fromkeys(motivos)) or "información insuficiente para validar avance"
    if missing_core >= 1:
        return "Moderado", ", ".join(dict.fromkeys(motivos)) or "información parcial; requiere verificación"
    return "Bajo", "desempeño académico estable"


def compare_risk(prev: Optional[str], current: str) -> str:
    if not prev or pd.isna(prev): return "Sin registro previo"
    if RIESGO_ORDEN.get(current, 0) < RIESGO_ORDEN.get(prev, 0): return "Mejora"
    if RIESGO_ORDEN.get(current, 0) > RIESGO_ORDEN.get(prev, 0): return "Empeora"
    return "Sin cambio"


def generate_message(row: pd.Series, asesor: str, horario: str, canal: str) -> str:
    nombre = row.get("nombre", "estudiante")
    riesgo = row.get("riesgo", "Bajo")
    motivo = row.get("motivo_detectado", "tu avance académico")
    if riesgo == "Bajo":
        return f"""Hola, {nombre}:\n\nEspero que te encontrés muy bien. Al revisar tu avance de esta semana, observo que mantenés un seguimiento adecuado del curso, especialmente en el cumplimiento de actividades, participación y desempeño general.\n\nTe felicito por la constancia que has demostrado y te animo a continuar con ese ritmo de trabajo. Recordá que cualquier duda o situación que necesités conversar podés escribirme por este medio.\n\nSaludos cordiales,\n{asesor}"""
    if riesgo == "Moderado":
        return f"""Hola, {nombre}:\n\nEspero que te encontrés bien. Al revisar tu avance académico de esta semana, identifiqué algunos aspectos que requieren atención, principalmente relacionados con {motivo}.\n\nEl propósito de este mensaje es brindarte acompañamiento oportuno y apoyarte para que puedas retomar el ritmo del curso antes de que la situación afecte de forma significativa tu desempeño académico. Te recomiendo revisar las actividades pendientes y organizar tus tiempos de entrega.\n\nQuedo atento a tu respuesta para conocer si existe alguna situación en la que podamos orientarte o apoyarte.\n\nSaludos cordiales,\n{asesor}"""
    return f"""Hola, {nombre}:\n\nEspero que te encontrés bien. Te escribo porque, al revisar tu participación y avance en Canvas, se identificaron señales importantes de riesgo académico relacionadas con {motivo}.\n\nEsta situación requiere atención prioritaria, ya que puede afectar tu continuidad y desempeño en el curso. Por ello, me gustaría brindarte un espacio de seguimiento más cercano.\n\nTe propongo que podamos reunirnos o comunicarnos en el siguiente horario: {horario}. Canal de atención: {canal}.\n\nQuedo atento a tu confirmación para poder apoyarte de manera oportuna.\n\nSaludos cordiales,\n{asesor}"""

# -----------------------------------------------------------------------------
# Derivation documents
# -----------------------------------------------------------------------------
def add_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text) if text is not None else "")
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def make_derivation_doc(row: pd.Series, asesor_academico: str, asesor_bienestar: str, observaciones: str, acciones: str) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx no está instalado.")
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(.55); sec.bottom_margin = Inches(.55); sec.left_margin = Inches(.65); sec.right_margin = Inches(.65)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = title.add_run("AVE UVG\nFormato de Derivación Académica")
    rr.bold = True; rr.font.size = Pt(16); rr.font.color.rgb = RGBColor(11, 31, 58)
    risk = row.get("riesgo", "")
    if risk == "Alto":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PRIORIDAD ALTA - RIESGO DE ABANDONO")
        run.bold = True; run.font.size = Pt(12); run.font.color.rgb = RGBColor(185, 28, 28)
    doc.add_paragraph(f"Fecha de derivación: {date.today().strftime('%d/%m/%Y')}")

    sections = [
        ("1. Datos del estudiante", [("Nombre del estudiante", row.get("nombre", "")), ("Carné", row.get("carne", "")), ("Carrera", row.get("carrera", "")), ("Trimestre", row.get("trimestre", "")), ("Correo electrónico", row.get("correo", "")), ("Teléfono", row.get("telefono", ""))]),
        ("2. Datos del remitente", [("Nombre del asesor académico", asesor_academico), ("Asesor de bienestar receptor", asesor_bienestar), ("Curso", row.get("curso", "")), ("Sección", row.get("seccion", ""))]),
        ("3. Motivo de la derivación", [("Nivel de riesgo", risk), ("Motivo detectado", row.get("motivo_detectado", ""))]),
        ("4. Descripción breve del caso", [("Descripción", f"El estudiante presenta nivel de riesgo {risk.lower()} debido a {row.get('motivo_detectado', 'indicadores académicos de alerta')}. Se recomienda seguimiento por parte del área de bienestar estudiantil conforme al protocolo institucional.")]),
        ("5. Acciones previas realizadas", [("Acciones", acciones)]),
        ("6. Observaciones adicionales", [("Observaciones", observaciones)]),
        ("7. Nivel de prioridad", [("Prioridad", "Alta (riesgo de abandono)" if risk == "Alto" else "Media (requiere apoyo pronto)")]),
    ]
    for head, items in sections:
        hp = doc.add_paragraph()
        run = hp.add_run(head)
        run.bold = True; run.font.color.rgb = RGBColor(11, 31, 58)
        table = doc.add_table(rows=len(items), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, (k, v) in enumerate(items):
            add_cell_text(table.cell(i,0), k, bold=True, color="#0B1F3A")
            add_cell_text(table.cell(i,1), v)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# -----------------------------------------------------------------------------
# Data processing
# -----------------------------------------------------------------------------
def standardize_analysis_input(df: pd.DataFrame, db: Dict[str, pd.DataFrame], curso: str, semana: int, asesor: str) -> pd.DataFrame:
    db = normalize_db_keys(db)
    df = normalize_key_columns(normalize_cols(df))
    aliases = {
        "nombre_del_estudiante": "nombre", "estudiante": "nombre", "sis_user_id": "carne", "id": "carne",
        "mail": "correo", "email": "correo", "score": "promedio", "calificacion": "promedio",
        "actividades_completadas_%": "actividades_pct", "porcentaje_actividades": "actividades_pct", "avance": "actividades_pct",
        "late": "entregas_tarde", "tardias": "entregas_tarde", "missing": "semanas_sin_entregas",
        "ingresos": "ingresos_semana", "dias_sin_actividad": "dias_inactivo", "respuesta_horas": "horas_respuesta"
    }
    df = df.rename(columns={c: aliases.get(c, c) for c in df.columns})
    for col in ["carne", "nombre", "correo", "telefono", "carrera", "trimestre", "curso", "seccion", "canvas_user_id", "asesor_bienestar", "actividades_pct", "promedio", "entregas_tarde", "semanas_sin_entregas", "ingresos_semana", "dias_inactivo", "horas_respuesta"]:
        if col not in df.columns:
            df[col] = np.nan
    df = normalize_key_columns(df)
    df["curso"] = df["curso"].fillna(curso)
    df.loc[df["curso"].astype(str).str.strip().eq(""), "curso"] = curso
    est = normalize_key_columns(db.get("Estudiantes", pd.DataFrame()))
    if not est.empty and "carne" in df.columns:
        df = normalize_key_columns(df)
        base_cols = [c for c in ["carne", "telefono", "carrera", "trimestre", "seccion", "asesor_bienestar", "canvas_user_id"] if c in est.columns]
        est_base = normalize_key_columns(est[base_cols].drop_duplicates("carne"))
        df = df.merge(est_base, on="carne", how="left", suffixes=("", "_base"))
        for c in ["telefono", "carrera", "trimestre", "seccion", "asesor_bienestar", "canvas_user_id"]:
            if f"{c}_base" in df.columns:
                df[c] = df[c].combine_first(df[f"{c}_base"])
                df.drop(columns=[f"{c}_base"], inplace=True)
    risks, motives = [], []
    for _, row in df.iterrows():
        r, m = classify_student(row)
        risks.append(r); motives.append(m)
    df["riesgo"] = risks
    df["motivo_detectado"] = motives
    hist = normalize_key_columns(db.get("Historial_Estudiantes", pd.DataFrame()))
    prev_map = {}
    if not hist.empty and "carne" in hist.columns:
        h = hist.dropna(subset=["carne"]).copy()
        h["fecha_sort"] = pd.to_datetime(h.get("fecha"), errors="coerce")
        h = h.sort_values("fecha_sort").drop_duplicates("carne", keep="last")
        prev_map = h.set_index("carne")["riesgo"].to_dict()
    df["riesgo_anterior"] = df["carne"].map(prev_map)
    df["cambio"] = [compare_risk(p, c) for p, c in zip(df["riesgo_anterior"], df["riesgo"])]
    df["semana"] = semana
    df["asesor_academico"] = asesor
    return df

# -----------------------------------------------------------------------------
# Excel local / sincronizado
# -----------------------------------------------------------------------------
def load_excel_db_from_path(path: str) -> Dict[str, pd.DataFrame]:
    """Lee una base Excel desde una ruta accesible para la app.

    Funciona muy bien cuando Streamlit se ejecuta en la computadora del asesor
    o en un servidor que tenga sincronizada una carpeta de OneDrive/SharePoint/Drive.
    """
    if not path or not str(path).strip():
        raise ValueError("Ingresá la ruta completa del archivo Excel.")
    path_obj = Path(str(path).strip().strip('\"'))
    if not path_obj.exists():
        raise FileNotFoundError(
            "No se encontró el archivo. Verificá que la ruta exista y que la carpeta sincronizada esté disponible en este equipo."
        )
    if path_obj.suffix.lower() not in [".xlsx", ".xlsm"]:
        raise ValueError("La base debe ser un archivo Excel .xlsx o .xlsm.")

    db = empty_db()
    xls = pd.ExcelFile(path_obj)
    for sheet in xls.sheet_names:
        key = next((s for s in SHEETS if s.lower() == sheet.lower()), sheet)
        if key in SHEETS:
            df = pd.read_excel(xls, sheet_name=sheet)
            df = normalize_cols(df)
            for col in SHEETS[key]:
                if col not in df.columns:
                    df[col] = np.nan
            db[key] = normalize_key_columns(df[SHEETS[key]])
    return normalize_db_keys(db)


def save_db_to_excel_path(db: Dict[str, pd.DataFrame], path: str, asesor: str = "asesor", make_backup: bool = True) -> Tuple[str, Optional[str]]:
    """Guarda la base actual en una ruta local/sincronizada y crea respaldo automático."""
    if not path or not str(path).strip():
        raise ValueError("Ingresá la ruta completa donde se guardará el Excel.")
    path_obj = Path(str(path).strip().strip('\"'))
    if path_obj.suffix.lower() not in [".xlsx", ".xlsm"]:
        raise ValueError("La ruta de guardado debe terminar en .xlsx o .xlsm.")
    path_obj.parent.mkdir(parents=True, exist_ok=True)

    backup_path = None
    if make_backup and path_obj.exists():
        safe_asesor = re.sub(r"[^A-Za-z0-9_-]+", "_", asesor or "asesor").strip("_")
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_dir = path_obj.parent / "backups_ave"
        backup_dir.mkdir(exist_ok=True)
        backup_path = backup_dir / f"backup_{path_obj.stem}_{stamp}_{safe_asesor}.xlsx"
        backup_path.write_bytes(path_obj.read_bytes())

    path_obj.write_bytes(export_db_excel(db))
    return str(path_obj), str(backup_path) if backup_path else None


def add_control_version(db: Dict[str, pd.DataFrame], asesor: str, curso: str, accion: str) -> Dict[str, pd.DataFrame]:
    """Registra en Configuracion una línea simple de control operativo.
    La hoja Configuracion sigue siendo de dos columnas para que sea fácil de entender.
    """
    registro = {
        "parametro": f"ultima_actualizacion_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
        "valor": f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | {accion} | asesor={asesor} | curso={curso}"
    }
    return append_rows(db, "Configuracion", [registro])

# -----------------------------------------------------------------------------
# Session state init
# -----------------------------------------------------------------------------
if "db" not in st.session_state:
    st.session_state.db = empty_db()
if "analysis_df" not in st.session_state:
    st.session_state.analysis_df = pd.DataFrame()
if "canvas_client" not in st.session_state:
    st.session_state.canvas_client = None
if "canvas_courses" not in st.session_state:
    st.session_state.canvas_courses = pd.DataFrame()
if "selected_canvas_course_id" not in st.session_state:
    st.session_state.selected_canvas_course_id = ""
if "selected_canvas_course_name" not in st.session_state:
    st.session_state.selected_canvas_course_name = ""
if "bienestar_df" not in st.session_state:
    st.session_state.bienestar_df = pd.DataFrame()
if "supabase_url" not in st.session_state:
    st.session_state.supabase_url = get_supabase_config().get("url", "")
if "supabase_key" not in st.session_state:
    st.session_state.supabase_key = get_supabase_config().get("service_role_key", get_supabase_config().get("anon_key", ""))
if "db_backend" not in st.session_state:
    st.session_state.db_backend = "Excel manual/local"

# -----------------------------------------------------------------------------
# Sidebar configuration
# -----------------------------------------------------------------------------
st.sidebar.image("https://dummyimage.com/480x100/0B1F3A/ffffff.png&text=AVE+Seguimiento+PRO", use_column_width=True)
st.sidebar.header("Configuración")
asesor_academico = st.sidebar.text_input("Nombre del asesor académico", value="Asesor Académico")
horario_atencion = st.sidebar.text_input("Horario de atención para riesgo alto", value="martes y jueves de 16:00 a 17:00")
canal_atencion = st.sidebar.text_input("Canal de atención", value="Bandeja de entrada de Canvas / enlace institucional")
semana_analisis = st.sidebar.selectbox("Semana de análisis", [1, 2, 3, 4, 5], index=0)
curso_manual = st.sidebar.text_input("Nombre del curso", value="Curso AVE")

st.title(APP_NAME)
st.caption("Clasificación de riesgo, historial académico, mensajes preventivos y derivaciones a bienestar con base de datos en Excel local/sincronizado o Supabase en línea.")

tabs = st.tabs(["🏠 Inicio", "🗂️ Base de datos", "🔌 Canvas / Datos", "📊 Dashboard", "👤 Estudiante", "✉️ Mensajes", "📌 Derivaciones", "⬇️ Exportar"])

# -----------------------------------------------------------------------------
# Inicio
# -----------------------------------------------------------------------------
with tabs[0]:
    st.markdown("### Propósito")
    st.write("Esta aplicación apoya el primer filtro del asesor académico mediante el análisis de actividades, calificaciones, entregas, actividad en Canvas y respuesta a comunicaciones. A partir de estos datos clasifica a cada estudiante en riesgo bajo, moderado o alto, genera acciones preventivas y registra historial para evitar duplicidad de derivaciones.")
    c1, c2, c3 = st.columns(3)
    c1.info("**Riesgo bajo**\n\nRetroalimentación cálida y seguimiento académico regular.")
    c2.warning("**Riesgo moderado**\n\nMensaje de preocupación, seguimiento y posible derivación a bienestar.")
    c3.error("**Riesgo alto**\n\nAtención prioritaria, horario de apoyo y derivación con distintivo de prioridad.")
    st.markdown("### Flujo recomendado")
    st.write("1. Cargar base de datos desde Excel o Supabase. 2. Conectar Canvas o cargar reporte. 3. Ejecutar análisis. 4. Revisar dashboard. 5. Generar mensajes. 6. Seleccionar derivaciones. 7. Guardar en Supabase o exportar Excel actualizado.")

# -----------------------------------------------------------------------------
# Base de datos
# -----------------------------------------------------------------------------
with tabs[1]:
    st.markdown("### Base de datos institucional")
    st.info(
        "La app puede trabajar con un Excel cargado manualmente, con un Excel en una carpeta sincronizada o con Supabase como base en línea para Streamlit Cloud. "
        "Para Streamlit Cloud, Supabase es el modo recomendado porque no depende de rutas locales ni Azure."
    )

    modo_bd = st.radio(
        "Método de trabajo",
        ["Cargar Excel manualmente", "Usar Excel desde ruta local / OneDrive / SharePoint sincronizado", "Usar Supabase en línea", "Crear base nueva"],
        horizontal=False,
    )

    if modo_bd == "Cargar Excel manualmente":
        uploaded_db = st.file_uploader("Cargar archivo Excel de base de datos", type=["xlsx"], key="db_upload")
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("Leer Excel cargado", type="primary"):
                try:
                    st.session_state.db = load_excel_db(uploaded_db)
                    st.session_state.db = add_control_version(st.session_state.db, asesor_academico, curso_manual, "Lectura de Excel cargado")
                    st.success("Base de datos cargada correctamente.")
                except Exception as e:
                    st.error(f"No se pudo leer el Excel: {e}")
        with col_b:
            st.download_button(
                "Descargar base actualizada",
                data=export_db_excel(st.session_state.db),
                file_name=f"base_datos_seguimiento_ave_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

    elif modo_bd == "Usar Excel desde ruta local / OneDrive / SharePoint sincronizado":
        st.markdown("#### Ruta del archivo sincronizado")
        st.caption("Ejemplo Windows: C:/Users/Usuario/Universidad del Valle de Guatemala/AVE/base_datos_seguimiento_ave.xlsx")
        excel_path = st.text_input("Ruta completa del Excel institucional", key="excel_sync_path")

        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("Probar acceso"):
                try:
                    path_obj = Path(excel_path.strip().strip('\"'))
                    if path_obj.exists():
                        st.success(f"Archivo encontrado: {path_obj.name}")
                        st.caption(f"Última modificación detectada: {datetime.fromtimestamp(path_obj.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')}")
                    else:
                        st.error("No se encontró el archivo en esa ruta.")
                except Exception as e:
                    st.error(f"No se pudo validar la ruta: {e}")
        with c2:
            if st.button("Leer base institucional", type="primary"):
                try:
                    st.session_state.db = load_excel_db_from_path(excel_path)
                    st.session_state.db = add_control_version(st.session_state.db, asesor_academico, curso_manual, "Lectura desde ruta sincronizada")
                    st.success("Base institucional leída correctamente.")
                except Exception as e:
                    st.error(f"No se pudo leer la base: {e}")
        with c3:
            if st.button("Guardar cambios en Excel"):
                try:
                    st.session_state.db = add_control_version(st.session_state.db, asesor_academico, curso_manual, "Guardado en ruta sincronizada")
                    saved_path, backup_path = save_db_to_excel_path(st.session_state.db, excel_path, asesor_academico, make_backup=True)
                    st.success(f"Base guardada correctamente en: {saved_path}")
                    if backup_path:
                        st.caption(f"Respaldo creado: {backup_path}")
                except Exception as e:
                    st.error(f"No se pudo guardar la base: {e}")

        st.warning(
            "Para evitar conflictos, no dejés el Excel abierto mientras la app guarda cambios. "
            "Si varios asesores usan la misma base, conviene guardar por turnos o trabajar con respaldos."
        )

    elif modo_bd == "Usar Supabase en línea":
        st.markdown("#### Conexión Supabase para Streamlit Cloud")
        st.caption("Recomendado: guardá la URL y la service_role_key en Streamlit Cloud > Settings > Secrets. La key no debe estar escrita en el código.")
        cfg = get_supabase_config()
        default_url = st.session_state.supabase_url or cfg.get("url", "")
        default_key = st.session_state.supabase_key or cfg.get("service_role_key", cfg.get("anon_key", ""))
        supa_url = st.text_input("SUPABASE_URL", value=default_url, key="supa_url_input")
        supa_key = st.text_input("SUPABASE_KEY", value=default_key, type="password", key="supa_key_input")
        st.session_state.supabase_url = supa_url
        st.session_state.supabase_key = supa_key

        st.code('[supabase]\nurl = "https://TU-PROYECTO.supabase.co"\nservice_role_key = "TU_SERVICE_ROLE_KEY"', language="toml")

        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("Probar Supabase"):
                try:
                    estado = supabase_status(supa_url, supa_key)
                    st.session_state.db_backend = "Supabase"
                    st.success("Conexión Supabase correcta.")
                    if estado.empty:
                        st.info("La tabla existe, pero todavía no hay hojas guardadas.")
                    else:
                        st.dataframe(estado, use_container_width=True, height=180)
                except Exception as e:
                    st.error(f"No se pudo conectar a Supabase: {e}")
                    st.info("Asegurate de crear la tabla con el archivo supabase_schema.sql incluido en el proyecto.")
        with c2:
            if st.button("Leer base desde Supabase", type="primary"):
                try:
                    st.session_state.db = read_db_from_supabase(supa_url, supa_key)
                    st.session_state.db = add_control_version(st.session_state.db, asesor_academico, curso_manual, "Lectura desde Supabase")
                    st.session_state.db_backend = "Supabase"
                    st.success("Base leída desde Supabase correctamente.")
                except Exception as e:
                    st.error(f"No se pudo leer la base desde Supabase: {e}")
        with c3:
            if st.button("Guardar base actual en Supabase"):
                try:
                    st.session_state.db = add_control_version(st.session_state.db, asesor_academico, curso_manual, "Guardado en Supabase")
                    save_db_to_supabase(st.session_state.db, supa_url, supa_key, asesor_academico)
                    st.session_state.db_backend = "Supabase"
                    st.success("Base guardada en Supabase correctamente.")
                except Exception as e:
                    st.error(f"No se pudo guardar en Supabase: {e}")

        st.warning("Usá la service_role_key únicamente en Streamlit Secrets. No la compartás con usuarios ni la subás a GitHub.")

    else:
        if st.button("Crear base vacía", type="primary"):
            st.session_state.db = empty_db()
            st.session_state.db = add_control_version(st.session_state.db, asesor_academico, curso_manual, "Creación de base nueva")
            st.success("Base vacía creada en memoria.")
        st.download_button(
            "Descargar plantilla Excel vacía",
            data=export_db_excel(empty_db()),
            file_name="base_datos_seguimiento_ave_template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    st.markdown("### Asignación de asesores de bienestar")
    st.caption("Cargá aquí el CSV o Excel institucional donde aparece el nombre del estudiante y su asesor de bienestar. La app hará match por carné y, si no existe, por nombre normalizado.")
    bienestar_file = st.file_uploader("Cargar base de bienestar CSV o Excel", type=["csv", "xlsx"], key="bienestar_upload")
    col_b1, col_b2 = st.columns([1, 2])
    with col_b1:
        if st.button("Aplicar asignación de bienestar"):
            try:
                bdf = read_bienestar_file(bienestar_file)
                st.session_state.bienestar_df = bdf
                st.session_state.db = apply_bienestar_assignment(st.session_state.db, bdf)
                if not st.session_state.analysis_df.empty:
                    st.session_state.analysis_df = apply_bienestar_to_analysis(st.session_state.analysis_df, st.session_state.db)
                st.success(f"Asignación aplicada. Asesores detectados: {bdf['asesor_bienestar'].nunique()} | Estudiantes en archivo: {len(bdf)}")
            except Exception as e:
                st.error(f"No se pudo aplicar la asignación de bienestar: {e}")
    with col_b2:
        if not st.session_state.bienestar_df.empty:
            resumen_b = st.session_state.bienestar_df.groupby("asesor_bienestar").size().reset_index(name="estudiantes_asignados").sort_values("estudiantes_asignados", ascending=False)
            st.dataframe(resumen_b, use_container_width=True, height=180)

    st.markdown("### Vista previa de la base actual")
    selected_sheet = st.selectbox("Hoja para revisar", list(SHEETS.keys()))
    preview_df = st.session_state.db.get(selected_sheet, pd.DataFrame()).copy()
    if selected_sheet == "Estudiantes" and (preview_df.empty or len(preview_df.dropna(how="all")) == 0):
        preview_df = latest_students_from_history(st.session_state.db)
        if not preview_df.empty:
            st.info("La hoja Estudiantes se reconstruyó automáticamente desde el último historial registrado.")
    st.dataframe(preview_df, use_container_width=True, height=320)

# -----------------------------------------------------------------------------
# Canvas / Datos
# -----------------------------------------------------------------------------
with tabs[2]:
    st.markdown("### Conexión con Canvas")
    url = st.text_input("URL de Canvas", placeholder="https://uvg.instructure.com")
    token = st.text_input("Token de Canvas", type="password")
    col1, col2, col3 = st.columns([1, 1, 2])
    with col1:
        if st.button("Validar token"):
            if not url or not token:
                st.warning("Ingresá la URL de Canvas y el token antes de validar.")
            else:
                try:
                    client = CanvasClient(url, token)
                    ok, msg = client.validate()
                    if ok:
                        st.session_state.canvas_client = client
                        st.success(f"Conexión validada: {msg}")
                    else:
                        st.error(msg)
                except Exception as e:
                    st.error(str(e))
    with col2:
        if st.button("Cargar cursos"):
            if not url or not token:
                st.warning("Ingresá la URL de Canvas y el token antes de cargar cursos.")
            else:
                try:
                    client = st.session_state.canvas_client or CanvasClient(url, token)
                    ok, msg = client.validate()
                    if not ok:
                        st.error(msg)
                    else:
                        st.session_state.canvas_client = client
                        cursos = client.courses()
                        st.session_state.canvas_courses = cursos
                        if cursos.empty:
                            st.warning("El token fue validado, pero Canvas no devolvió cursos activos para este usuario.")
                        else:
                            st.success(f"Se encontraron {len(cursos)} cursos disponibles.")
                except Exception as e:
                    st.error(f"No se pudieron cargar los cursos: {e}")
    with col3:
        st.info("Primero valide el token y luego cargue los cursos. La lista mostrará únicamente los cursos a los que el token tenga acceso.")

    cursos_df = st.session_state.canvas_courses.copy()
    selected_course_id = st.session_state.selected_canvas_course_id
    selected_course_name = st.session_state.selected_canvas_course_name or curso_manual
    if not cursos_df.empty:
        labels = cursos_df["label"].tolist() if "label" in cursos_df.columns else cursos_df["name"].astype(str).tolist()
        default_index = 0
        if selected_course_id:
            match = cursos_df.index[cursos_df["id"].astype(str).eq(str(selected_course_id))].tolist()
            if match:
                default_index = int(match[0])
        selected_label = st.selectbox("Seleccione el curso de Canvas a analizar", labels, index=default_index)
        selected_row = cursos_df.loc[(cursos_df["label"] if "label" in cursos_df.columns else cursos_df["name"].astype(str)).eq(selected_label)].iloc[0]
        selected_course_id = str(selected_row["id"])
        selected_course_name = str(selected_row["name"])
        st.session_state.selected_canvas_course_id = selected_course_id
        st.session_state.selected_canvas_course_name = selected_course_name
        st.caption(f"Curso seleccionado: {selected_course_name} | ID Canvas: {selected_course_id}")
        st.dataframe(cursos_df.drop(columns=["label"], errors="ignore"), use_container_width=True, height=180)
    else:
        st.warning("Aún no hay cursos cargados. Usá el botón **Cargar cursos** para obtenerlos desde Canvas.")

    st.markdown("### Cargar reporte manual")
    report_file = st.file_uploader("Cargar reporte CSV o Excel con columnas de indicadores", type=["csv", "xlsx"], key="report")
    if report_file:
        if report_file.name.lower().endswith("csv"):
            raw = pd.read_csv(report_file)
        else:
            raw = pd.read_excel(report_file)
        st.write("Vista previa del reporte cargado")
        st.dataframe(raw.head(20), use_container_width=True)
        if st.button("Ejecutar análisis con reporte cargado", type="primary"):
            curso_analisis = st.session_state.selected_canvas_course_name or curso_manual
            analyzed = standardize_analysis_input(raw, st.session_state.db, curso_analisis, semana_analisis, asesor_academico)
            analyzed = apply_bienestar_to_analysis(analyzed, st.session_state.db)
            st.session_state.analysis_df = analyzed
            st.session_state.db = upsert_students_from_analysis(st.session_state.db, analyzed)
            idc = datetime.now().strftime("C%Y%m%d%H%M%S")
            rows = []
            for _, r in analyzed.iterrows():
                d = r.to_dict(); d["fecha"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S"); d["id_consulta"] = idc
                rows.append(d)
            append_rows(st.session_state.db, "Historial_Estudiantes", rows)
            counts = analyzed["riesgo"].value_counts().to_dict()
            append_rows(st.session_state.db, "Consultas_Canvas", [{"id_consulta": idc, "fecha_consulta": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "asesor_academico": asesor_academico, "curso": (st.session_state.selected_canvas_course_name or curso_manual), "semana": semana_analisis, "total_estudiantes": len(analyzed), "bajo": counts.get("Bajo", 0), "moderado": counts.get("Moderado", 0), "alto": counts.get("Alto", 0), "fuente_datos": "reporte manual"}])
            st.success("Análisis ejecutado y registrado en historial.")

    with st.expander("Extracción básica desde Canvas", expanded=True):
        course_id = st.session_state.selected_canvas_course_id
        course_name = st.session_state.selected_canvas_course_name or curso_manual
        if course_id:
            st.success(f"Curso listo para consultar: {course_name} | ID: {course_id}")
        else:
            st.warning("Primero cargá y seleccioná un curso de Canvas. También podés escribir el ID manualmente si ya lo conocés.")
            course_id = st.text_input("ID del curso Canvas", value="")
            course_name = curso_manual

        if st.button("Obtener estudiantes del curso seleccionado"):
            if not course_id:
                st.warning("Seleccioná un curso o ingresá manualmente el ID del curso Canvas.")
            else:
                try:
                    client = st.session_state.canvas_client or CanvasClient(url, token)
                    users = client.course_metrics(course_id, course_name)
                    analyzed = standardize_analysis_input(users, st.session_state.db, course_name, semana_analisis, asesor_academico)
                    analyzed = apply_bienestar_to_analysis(analyzed, st.session_state.db)
                    st.session_state.analysis_df = analyzed
                    st.session_state.db = upsert_students_from_analysis(st.session_state.db, analyzed)
                    idc = datetime.now().strftime("C%Y%m%d%H%M%S")
                    rows = []
                    for _, r in analyzed.iterrows():
                        d = r.to_dict(); d["fecha"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S"); d["id_consulta"] = idc
                        rows.append(d)
                    append_rows(st.session_state.db, "Historial_Estudiantes", rows)
                    counts = analyzed["riesgo"].value_counts().to_dict()
                    append_rows(st.session_state.db, "Consultas_Canvas", [{"id_consulta": idc, "fecha_consulta": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "asesor_academico": asesor_academico, "curso": course_name, "semana": semana_analisis, "total_estudiantes": len(analyzed), "bajo": counts.get("Bajo", 0), "moderado": counts.get("Moderado", 0), "alto": counts.get("Alto", 0), "fuente_datos": "Canvas API"}])
                    st.success("Análisis ejecutado desde Canvas y registrado en historial. Si Canvas no expone algún dato, la app lo toma como alerta para evitar falsos riesgos bajos.")
                    st.dataframe(st.session_state.analysis_df, use_container_width=True)
                except Exception as e:
                    st.error(f"No se pudo obtener información: {e}")

# -----------------------------------------------------------------------------
# Dashboard
# -----------------------------------------------------------------------------
with tabs[3]:
    df = st.session_state.analysis_df.copy()
    if df.empty:
        st.warning("Primero ejecutá un análisis desde la pestaña Canvas / Datos.")
    else:
        f1, f2 = st.columns(2)
        risk_filter = f1.multiselect("Filtrar por riesgo", ["Bajo", "Moderado", "Alto"], default=["Bajo", "Moderado", "Alto"])
        section_filter = f2.multiselect("Filtrar por sección", sorted(df["seccion"].dropna().astype(str).unique()))
        dff = df[df["riesgo"].isin(risk_filter)]
        if section_filter:
            dff = dff[dff["seccion"].astype(str).isin(section_filter)]
        total = len(dff)
        bajo = int((dff["riesgo"] == "Bajo").sum()); mod = int((dff["riesgo"] == "Moderado").sum()); alto = int((dff["riesgo"] == "Alto").sum())
        k1, k2, k3, k4 = st.columns(4)
        k1.metric("Estudiantes analizados", total)
        k2.metric("Riesgo bajo", bajo)
        k3.metric("Riesgo moderado", mod)
        k4.metric("Riesgo alto", alto)
        cc1, cc2 = st.columns([1, 2])
        with cc1:
            fig = px.pie(dff, names="riesgo", title="Distribución por riesgo", color="riesgo", color_discrete_map=RIESGO_COLOR)
            st.plotly_chart(fig, use_container_width=True)
        with cc2:
            if "cambio" in dff.columns:
                fig2 = px.histogram(dff, x="cambio", color="riesgo", title="Evolución respecto al registro anterior", color_discrete_map=RIESGO_COLOR)
                st.plotly_chart(fig2, use_container_width=True)
        st.markdown("### Tabla de seguimiento")
        cols = [c for c in ["carne", "nombre", "correo", "curso", "seccion", "actividades_pct", "promedio", "entregas_tarde", "ingresos_semana", "dias_inactivo", "horas_respuesta", "riesgo", "riesgo_anterior", "cambio", "motivo_detectado", "asesor_bienestar"] if c in dff.columns]
        st.dataframe(dff[cols], use_container_width=True, height=420)

# -----------------------------------------------------------------------------
# Estudiante
# -----------------------------------------------------------------------------
with tabs[4]:
    df = st.session_state.analysis_df.copy()
    if df.empty:
        st.warning("No hay análisis activo.")
    else:
        names = (df["nombre"].fillna("") + " | " + df["carne"].fillna("").astype(str)).tolist()
        idx = st.selectbox("Seleccionar estudiante", range(len(names)), format_func=lambda i: names[i])
        row = df.iloc[idx]
        c1, c2, c3 = st.columns(3)
        c1.metric("Riesgo actual", row.get("riesgo", ""))
        c2.metric("Riesgo anterior", row.get("riesgo_anterior", "Sin registro"))
        c3.metric("Cambio", row.get("cambio", ""))
        st.write("**Motivo detectado:**", row.get("motivo_detectado", ""))
        st.dataframe(pd.DataFrame([row]), use_container_width=True)
        hist = st.session_state.db.get("Historial_Estudiantes", pd.DataFrame())
        if not hist.empty and row.get("carne") in hist.get("carne", pd.Series()).astype(str).values:
            st.markdown("### Historial registrado")
            st.dataframe(hist[hist["carne"].astype(str) == str(row.get("carne"))].tail(10), use_container_width=True)

# -----------------------------------------------------------------------------
# Mensajes
# -----------------------------------------------------------------------------
with tabs[5]:
    df = st.session_state.analysis_df.copy()
    if df.empty:
        st.warning("No hay análisis activo.")
    else:
        risk_msg = st.multiselect("Generar mensajes para riesgo", ["Bajo", "Moderado", "Alto"], default=["Moderado", "Alto"])
        candidates = df[df["riesgo"].isin(risk_msg)].copy()
        selected = st.multiselect("Seleccionar estudiantes", candidates.index.tolist(), format_func=lambda i: f"{candidates.loc[i,'nombre']} | {candidates.loc[i,'riesgo']}")
        if selected:
            for i in selected:
                row = candidates.loc[i]
                msg = generate_message(row, asesor_academico, horario_atencion, canal_atencion)
                with st.expander(f"{row.get('nombre')} - {row.get('riesgo')}"):
                    edited = st.text_area("Mensaje", msg, height=230, key=f"msg_{i}")
                    colx, coly = st.columns(2)
                    if colx.button("Registrar mensaje", key=f"regmsg_{i}"):
                        append_rows(st.session_state.db, "Mensajes_Enviados", [{"fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "carne": row.get("carne"), "nombre": row.get("nombre"), "correo": row.get("correo"), "curso": row.get("curso"), "riesgo": row.get("riesgo"), "tipo_mensaje": f"Seguimiento {row.get('riesgo')}", "mensaje_generado": edited, "enviado_canvas": "No", "asesor_academico": asesor_academico}])
                        st.success("Mensaje registrado en la base.")
                    if coly.button("Enviar por Canvas", key=f"send_{i}"):
                        client = st.session_state.canvas_client or (CanvasClient(url, token) if url and token else None)
                        if not client:
                            st.error("Primero valide Canvas o ingrese URL/token.")
                        else:
                            recip = str(row.get("canvas_user_id") or row.get("correo") or "").strip()
                            ok, resp = client.send_message([recip], f"Seguimiento académico - {row.get('curso')}", edited)
                            st.success(resp) if ok else st.error(resp)

# -----------------------------------------------------------------------------
# Derivaciones
# -----------------------------------------------------------------------------
with tabs[6]:
    df = st.session_state.analysis_df.copy()
    if df.empty:
        st.warning("No hay análisis activo.")
    else:
        df = apply_bienestar_to_analysis(df, st.session_state.db)
        st.session_state.analysis_df = df
        derivables = df[df["riesgo"].isin(["Moderado", "Alto"])].copy()
        if derivables.empty:
            st.success("No hay estudiantes en riesgo moderado o alto para derivar en el análisis activo.")
        else:
            st.markdown("### Filtro por asesor de bienestar")
            if "asesor_bienestar" not in derivables.columns:
                derivables["asesor_bienestar"] = "Sin asesor asignado"
            derivables["asesor_bienestar"] = derivables["asesor_bienestar"].fillna("Sin asesor asignado").replace("", "Sin asesor asignado")
            asesores_detectados = sorted(derivables["asesor_bienestar"].dropna().astype(str).unique().tolist())
            asesor_filtro = st.selectbox("Mostrar derivaciones que corresponden a", ["Todos"] + asesores_detectados)
            if asesor_filtro != "Todos":
                derivables = derivables[derivables["asesor_bienestar"].astype(str).eq(asesor_filtro)].copy()

            c1, c2, c3 = st.columns(3)
            c1.metric("Derivables mostrados", len(derivables))
            c2.metric("Riesgo moderado", int((derivables["riesgo"] == "Moderado").sum()))
            c3.metric("Riesgo alto", int((derivables["riesgo"] == "Alto").sum()))

            st.dataframe(
                derivables[[c for c in ["carne", "nombre", "correo", "curso", "seccion", "riesgo", "cambio", "asesor_bienestar", "motivo_detectado"] if c in derivables.columns]],
                use_container_width=True,
                height=280,
            )

            obs = st.text_area("Observaciones adicionales para los formatos", height=100)
            acciones = st.text_area(
                "Acciones previas realizadas",
                value="Se revisó avance académico en Canvas, se identificó el nivel de riesgo, se generó mensaje de seguimiento y se registra el caso para acompañamiento oportuno.",
                height=100,
            )
            selected = st.multiselect(
                "Seleccionar estudiantes para derivar",
                derivables.index.tolist(),
                format_func=lambda i: f"{derivables.loc[i,'nombre']} | {derivables.loc[i,'riesgo']} | Bienestar: {derivables.loc[i].get('asesor_bienestar','Sin asesor')}",
            )
            if selected and st.button("Generar paquete de derivación", type="primary"):
                zip_buffer = io.BytesIO()
                report_rows = []
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i in selected:
                        row = derivables.loc[i].copy()
                        advisor = str(row.get("asesor_bienestar") or "Sin asesor asignado")
                        row["observaciones"] = obs
                        doc_bytes = make_derivation_doc(row, asesor_academico, advisor, obs, acciones)
                        safe_student = re.sub(r"[^A-Za-z0-9_-]+", "_", str(row.get("nombre", "estudiante")))[:45]
                        safe_advisor = re.sub(r"[^A-Za-z0-9_-]+", "_", advisor).strip("_") or "Sin_asesor"
                        zf.writestr(f"{safe_advisor}/derivacion_{safe_student}_{row.get('riesgo')}.docx", doc_bytes)
                        prioridad = "Alta" if row.get("riesgo") == "Alto" else "Media"
                        report_rows.append({
                            "id_derivacion": datetime.now().strftime("D%Y%m%d%H%M%S") + str(i),
                            "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            "carne": row.get("carne"),
                            "nombre": row.get("nombre"),
                            "correo": row.get("correo"),
                            "curso": row.get("curso"),
                            "seccion": row.get("seccion"),
                            "riesgo": row.get("riesgo"),
                            "prioridad": prioridad,
                            "asesor_bienestar": advisor,
                            "correo_bienestar": "",
                            "motivo": row.get("motivo_detectado"),
                            "acciones_previas": acciones,
                            "observaciones": obs,
                            "estado_derivacion": "Generada",
                            "asesor_academico": asesor_academico,
                        })
                    report_df = pd.DataFrame(report_rows)
                    bio = io.BytesIO()
                    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
                        report_df.to_excel(writer, index=False, sheet_name="Listado_Derivaciones")
                        resumen = report_df.groupby(["asesor_bienestar", "riesgo"]).size().reset_index(name="cantidad") if not report_df.empty else pd.DataFrame()
                        resumen.to_excel(writer, index=False, sheet_name="Resumen_por_Asesor")
                    zf.writestr("listado_general_derivaciones.xlsx", bio.getvalue())
                append_rows(st.session_state.db, "Derivaciones", report_rows)
                st.success("Paquete generado y derivaciones registradas. El ZIP queda organizado por carpetas de asesor de bienestar.")
                st.download_button("Descargar paquete ZIP", zip_buffer.getvalue(), file_name=f"paquete_derivaciones_{date.today()}.zip", mime="application/zip")

# -----------------------------------------------------------------------------
# Export
# -----------------------------------------------------------------------------
with tabs[7]:
    st.markdown("### Exportar base actualizada")
    if (st.session_state.db.get("Estudiantes", pd.DataFrame()).empty) and not st.session_state.db.get("Historial_Estudiantes", pd.DataFrame()).empty:
        st.session_state.db["Estudiantes"] = latest_students_from_history(st.session_state.db)
    bytes_xlsx = export_db_excel(st.session_state.db)
    est_count = len(st.session_state.db.get("Estudiantes", pd.DataFrame()).dropna(how="all"))
    hist_count = len(st.session_state.db.get("Historial_Estudiantes", pd.DataFrame()).dropna(how="all"))
    st.caption(f"Registros a exportar: Estudiantes={est_count} | Historial={hist_count} | Backend activo={st.session_state.get('db_backend', 'Excel')}")
    col_exp1, col_exp2 = st.columns(2)
    with col_exp1:
        st.download_button("Descargar base de datos actualizada Excel", bytes_xlsx, file_name="base_datos_seguimiento_ave_actualizada.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with col_exp2:
        if st.button("Guardar cambios en Supabase"):
            try:
                save_db_to_supabase(st.session_state.db, st.session_state.supabase_url, st.session_state.supabase_key, asesor_academico)
                st.session_state.db_backend = "Supabase"
                st.success("Cambios guardados en Supabase.")
            except Exception as e:
                st.error(f"No se pudo guardar en Supabase: {e}")
    if not st.session_state.analysis_df.empty:
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
            st.session_state.analysis_df.to_excel(writer, index=False, sheet_name="Analisis_Actual")
        st.download_button("Descargar análisis actual", out.getvalue(), file_name="analisis_actual_estudiantes.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
